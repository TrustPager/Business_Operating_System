#!/usr/bin/env python3
"""write_docx.py — the doc-lib-set WRITE side for Word (.docx).

The keyless WRITE counterpart to markitdown_convert.py (the READ side). Turns a
small structured document spec Claude produced into a real .docx an owner can
open and send — proposals, policies, job ads, letters — no account, no network,
one local install.

Usage:
    # blocks as JSON on the command line:
    python tools/write_docx.py --out proposal.docx --blocks \
        '[{"type":"heading","text":"Proposal","level":1},
          {"type":"paragraph","text":"Thanks for the opportunity."}]'

    # or pipe the JSON in on stdin:
    cat doc.json | python tools/write_docx.py --out proposal.docx

Input shape: a JSON array of blocks. Each block is an object with a "type":
    {"type":"heading","text":"...","level":1}   # level 0-9, default 1
    {"type":"paragraph","text":"..."}            # a body paragraph
    {"type":"bullet","text":"..."}               # a bulleted list item
    {"type":"table","header":["Item","Qty","Price"],   # header row (optional)
     "rows":[["Site visit","1","$120"],              # one list per data row
             ["Labour","2","$180"]]}                 # a real priced line-item TABLE

The table block renders a real Word table (header row bolded when present),
so a quote or proposal lays its priced line items out in a grid the owner can
send, instead of bullets with `$____`.

UTF-8 / Windows note: when the blocks JSON carries any non-ASCII characters
(en-dashes, curly quotes, accented names, emoji), DO NOT inline it after
--blocks on the Windows command line. The console code page (often cp1252)
mangles those into mojibake before Python sees them, so they land wrong in the
document. Write the JSON to a UTF-8 temp file and pipe it in on stdin instead:
    python tools/write_docx.py --out proposal.docx < blocks.json
Stdin is read as UTF-8 explicitly here, so a UTF-8 temp file round-trips
cleanly. Plain-ASCII payloads are fine to inline.

This is part of the doc-lib-set keyless WRITE driver (`doclib`). See
knowledge/document-tools-method.md. No network at runtime.
"""
import sys
import json
import argparse


# The pip spec the BOS installs when this lib is missing. python-docx imports
# as `docx`, so the missing-import branch keys off `docx` but installs `python-docx`.
MISSING_DEP_SPEC = "python-docx"

# Machine-readable + human missing-dependency signal (D11). The leading
# BOS_MISSING_DEP: line is what the SKILL layer keys off to run a detect ->
# offer -> install-on-yes -> verify loop; the second line recommends
# `python -m pip install` (never bare `pip`, a churn trap on multi-Python Windows).
INSTALL_HINT = (
    f"BOS_MISSING_DEP: {MISSING_DEP_SPEC}\n"
    "python-docx isn't installed (the doc-lib-set tool for writing .docx files).\n"
    f"Install it with: python -m pip install {MISSING_DEP_SPEC}\n"
)


def _read_stdin_utf8() -> str:
    """Read stdin as UTF-8 regardless of the console code page.

    On Windows ``sys.stdin.read()`` decodes using the locale code page (often
    cp1252), which mangles en-dashes / curly quotes / accents into mojibake. The
    UTF-8-safe arg path is a UTF-8 temp file piped in on stdin, so decode the raw
    bytes as UTF-8 here. Falls back to the text stream if the buffer is absent.
    """
    buf = getattr(sys.stdin, "buffer", None)
    if buf is not None:
        return buf.read().decode("utf-8")
    return sys.stdin.read()


def _load_blocks(blocks_arg: str | None) -> list[dict]:
    """Load the blocks JSON from --blocks or stdin; validate it is a list of objects."""
    if blocks_arg is not None:
        raw = blocks_arg
    else:
        raw = _read_stdin_utf8()
    if not raw.strip():
        sys.stderr.write("No blocks provided (pass --blocks '<json>' or pipe JSON on stdin).\n")
        sys.exit(1)
    try:
        blocks = json.loads(raw)
    except json.JSONDecodeError as e:
        sys.stderr.write(f"Could not parse blocks JSON: {e}\n")
        sys.exit(1)
    if not isinstance(blocks, list) or not all(isinstance(b, dict) for b in blocks):
        sys.stderr.write(
            "Blocks must be a JSON array of objects, e.g. "
            "[{\"type\":\"paragraph\",\"text\":\"...\"}].\n"
        )
        sys.exit(1)
    return blocks


def _add_table(doc, block: dict) -> None:
    """Render a {"type":"table", "header":[...], "rows":[[...], ...]} block.

    A real Word table so priced line items land in a grid, not bullets with
    `$____`. ``header`` is an optional list of column labels (rendered bold);
    ``rows`` is a list of rows, each a list of cell values. Every cell is
    coerced to a string. The column count is taken from the header if present,
    else from the widest row. Short rows are padded, long rows truncated, so a
    ragged payload still produces a valid table rather than crashing.
    """
    header = block.get("header")
    rows = block.get("rows", [])
    if header is not None and not isinstance(header, list):
        sys.stderr.write("table block: 'header' must be a list of column labels.\n")
        sys.exit(1)
    if not isinstance(rows, list) or not all(isinstance(r, list) for r in rows):
        sys.stderr.write("table block: 'rows' must be a list of rows (each a list of cells).\n")
        sys.exit(1)
    if header is None and not rows:
        sys.stderr.write("table block: needs a 'header' or at least one row.\n")
        sys.exit(1)

    ncols = len(header) if header else max((len(r) for r in rows), default=0)
    if ncols == 0:
        sys.stderr.write("table block: could not determine any columns.\n")
        sys.exit(1)

    def _cells(values: list) -> list[str]:
        cells = [("" if v is None else str(v)) for v in values][:ncols]
        cells += [""] * (ncols - len(cells))
        return cells

    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    if header:
        hdr = table.add_row().cells
        for cell, value in zip(hdr, _cells(header)):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, _cells(row)):
            cell.text = value


def write_docx(path: str, blocks: list[dict]) -> None:
    try:
        from docx import Document
    except ImportError:
        sys.stderr.write(INSTALL_HINT)
        sys.exit(2)

    doc = Document()
    try:
        for block in blocks:
            btype = block.get("type", "paragraph")
            text = block.get("text", "")
            if btype == "heading":
                level = block.get("level", 1)
                try:
                    level = int(level)
                except (TypeError, ValueError):
                    level = 1
                level = max(0, min(level, 9))
                doc.add_heading(text, level=level)
            elif btype == "bullet":
                doc.add_paragraph(text, style="List Bullet")
            elif btype == "paragraph":
                doc.add_paragraph(text)
            elif btype == "table":
                _add_table(doc, block)
            else:
                sys.stderr.write(
                    f"Unknown block type {btype!r} (expected heading/paragraph/bullet/table).\n"
                )
                sys.exit(1)
        doc.save(path)
    except SystemExit:
        raise
    except Exception as e:  # noqa: BLE001 — surface the real reason to the operator
        sys.stderr.write(f"Could not write {path}: {e}\n")
        sys.exit(1)


def main() -> None:
    ap = argparse.ArgumentParser(description="Write a real .docx from JSON blocks (python-docx).")
    ap.add_argument("--out", required=True, help="Path to write the .docx file to.")
    ap.add_argument("--blocks", help="Blocks as a JSON array of objects. Omit to read JSON from stdin.")
    args = ap.parse_args()

    blocks = _load_blocks(args.blocks)
    write_docx(args.out, blocks)
    sys.stderr.write(f"Wrote {len(blocks)} block(s) to {args.out}\n")


if __name__ == "__main__":
    main()
