"""Tests for the doc-lib-set keyless WRITE driver (Floor Wave 1, Task 1.1).

The doc-lib-set is the keyless WRITE counterpart to markitdown_convert.py (the
READ side). Four thin tools/ wrappers, each mirroring the MarkItDown wrapper's
shape (argparse CLI, INSTALL_HINT, exit 2 when the lib is missing, exit 1 on
error):

    tools/write_xlsx.py   — write .xlsx        (openpyxl)
    tools/write_docx.py   — write .docx        (python-docx)
    tools/make_pdf.py     — generate a PDF     (reportlab)
    tools/pdf_tables.py   — extract PDF tables (pdfplumber)

Two kinds of test here:

  1. Round-trips (write then read back) and the table extraction. These need the
     vendored libs installed, so each is guarded with @skipUnless(<lib>). The
     offline suite stays GREEN on a machine without the libs — the round-trips
     simply skip. Install them (pip install openpyxl python-docx pdfplumber
     reportlab) to actually exercise the round-trips.

  2. The missing-lib install-hint path (exit 2). This is tested UNCONDITIONALLY:
     we run each wrapper as a subprocess with the relevant import blocked, and
     assert it exits 2 and prints a one-line `pip install` hint to stderr. This
     never needs the lib, so it always runs.

Offline-safe: no network, no key. Run:
    BOS_OFFLINE=1 python -m unittest tests.test_doc_lib_set
"""

import json
import subprocess
import sys
import tempfile
import unittest
from importlib import util as importutil
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
TOOLS = REPO / "tools"


def _lib(name: str) -> bool:
    """True if importable (the round-trip can run)."""
    return importutil.find_spec(name) is not None


HAS_OPENPYXL = _lib("openpyxl")
HAS_DOCX = _lib("docx")
HAS_REPORTLAB = _lib("reportlab")
HAS_PDFPLUMBER = _lib("pdfplumber")


def _run_wrapper(script: str, args: list[str], *, block_import: str | None = None,
                 stdin: str | None = None) -> subprocess.CompletedProcess:
    """Run a tools/ wrapper as a subprocess.

    When ``block_import`` is set, a tiny sitecustomize-style preamble installs a
    meta-path finder that makes ``import <block_import>`` raise ImportError, so
    we can deterministically exercise the wrapper's missing-lib branch even when
    the lib IS installed in this environment.
    """
    if block_import is None:
        cmd = [sys.executable, str(TOOLS / script), *args]
        return subprocess.run(cmd, capture_output=True, text=True, input=stdin)

    # Block the target import via a meta-path finder, then exec the real script.
    preamble = (
        "import sys\n"
        f"_BLOCK = {block_import!r}\n"
        "class _Blocker:\n"
        "    def find_spec(self, name, path=None, target=None):\n"
        "        if name == _BLOCK or name.startswith(_BLOCK + '.'):\n"
        "            raise ImportError('blocked for test: ' + name)\n"
        "        return None\n"
        "sys.meta_path.insert(0, _Blocker())\n"
        f"sys.argv = [{str(TOOLS / script)!r}, *{args!r}]\n"
        f"exec(compile(open({str(TOOLS / script)!r}, encoding='utf-8').read(), "
        f"{str(TOOLS / script)!r}, 'exec'))\n"
    )
    return subprocess.run([sys.executable, "-c", preamble],
                          capture_output=True, text=True, input=stdin)


class TestInstallHintPaths(unittest.TestCase):
    """The missing-lib path: exit 2 + a one-line pip install hint. Always runs."""

    def test_write_xlsx_missing_lib_exits_2_with_hint(self):
        proc = _run_wrapper("write_xlsx.py", ["--out", "x.xlsx", "--rows", "[[1]]"],
                            block_import="openpyxl")
        self.assertEqual(proc.returncode, 2, proc.stderr)
        self.assertIn("pip install openpyxl", proc.stderr)

    def test_write_docx_missing_lib_exits_2_with_hint(self):
        proc = _run_wrapper("write_docx.py",
                            ["--out", "x.docx", "--blocks", '[{"type":"paragraph","text":"hi"}]'],
                            block_import="docx")
        self.assertEqual(proc.returncode, 2, proc.stderr)
        self.assertIn("pip install python-docx", proc.stderr)

    def test_make_pdf_missing_lib_exits_2_with_hint(self):
        proc = _run_wrapper("make_pdf.py",
                            ["--out", "x.pdf", "--blocks", '[{"type":"paragraph","text":"hi"}]'],
                            block_import="reportlab")
        self.assertEqual(proc.returncode, 2, proc.stderr)
        self.assertIn("pip install reportlab", proc.stderr)

    def test_pdf_tables_missing_lib_exits_2_with_hint(self):
        proc = _run_wrapper("pdf_tables.py", ["nonexistent.pdf"],
                            block_import="pdfplumber")
        self.assertEqual(proc.returncode, 2, proc.stderr)
        self.assertIn("pip install pdfplumber", proc.stderr)


class TestErrorPaths(unittest.TestCase):
    """Bad input exits 1 (not 2, not a traceback). Independent of the libs."""

    def test_write_xlsx_bad_json_exits_1(self):
        proc = _run_wrapper("write_xlsx.py", ["--out", "x.xlsx", "--rows", "not-json"])
        self.assertEqual(proc.returncode, 1, proc.stderr)

    def test_write_docx_bad_shape_exits_1(self):
        # valid JSON, wrong shape (not a list of objects)
        proc = _run_wrapper("write_docx.py", ["--out", "x.docx", "--blocks", '["just a string"]'])
        self.assertEqual(proc.returncode, 1, proc.stderr)


@unittest.skipUnless(HAS_OPENPYXL, "openpyxl not installed")
class TestXlsxRoundTrip(unittest.TestCase):
    def test_write_then_read_back_xlsx(self):
        from openpyxl import load_workbook
        with tempfile.TemporaryDirectory() as d:
            out = Path(d) / "rt.xlsx"
            rows = [["Item", "Qty", "Price"], ["Site visit", 1, 120], ["Labour", 2, 90]]
            proc = _run_wrapper(
                "write_xlsx.py",
                ["--out", str(out), "--rows", json.dumps(rows), "--header", "--sheet", "Quote"],
            )
            self.assertEqual(proc.returncode, 0, proc.stderr)
            self.assertTrue(out.is_file())

            wb = load_workbook(out)
            ws = wb.active
            self.assertEqual(ws.title, "Quote")
            read = [[c.value for c in row] for row in ws.iter_rows()]
            self.assertEqual(read, rows)
            # header row rendered bold
            self.assertTrue(ws.cell(row=1, column=1).font.bold)


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class TestDocxRoundTrip(unittest.TestCase):
    def test_write_then_read_back_docx(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as d:
            out = Path(d) / "rt.docx"
            blocks = [
                {"type": "heading", "text": "Proposal", "level": 1},
                {"type": "paragraph", "text": "Thanks for the opportunity."},
                {"type": "bullet", "text": "Scope confirmed."},
            ]
            proc = _run_wrapper(
                "write_docx.py", ["--out", str(out), "--blocks", json.dumps(blocks)]
            )
            self.assertEqual(proc.returncode, 0, proc.stderr)
            self.assertTrue(out.is_file())

            doc = Document(str(out))
            texts = [p.text for p in doc.paragraphs]
            self.assertIn("Proposal", texts)
            self.assertIn("Thanks for the opportunity.", texts)
            self.assertIn("Scope confirmed.", texts)

    def test_write_then_read_back_table_block(self):
        """A table block renders a real Word table (priced line items as a grid),
        with the header row bold, round-tripping every cell back out."""
        from docx import Document
        with tempfile.TemporaryDirectory() as d:
            out = Path(d) / "rt_table.docx"
            blocks = [
                {"type": "heading", "text": "Your investment", "level": 2},
                {
                    "type": "table",
                    "header": ["Item", "Qty", "Price"],
                    "rows": [
                        ["Site visit", "1", "$120"],
                        ["Labour", "2", "$180"],
                    ],
                },
            ]
            proc = _run_wrapper(
                "write_docx.py", ["--out", str(out), "--blocks", json.dumps(blocks)]
            )
            self.assertEqual(proc.returncode, 0, proc.stderr)
            self.assertTrue(out.is_file())

            doc = Document(str(out))
            self.assertEqual(len(doc.tables), 1)
            table = doc.tables[0]
            grid = [[c.text for c in row.cells] for row in table.rows]
            self.assertEqual(
                grid,
                [
                    ["Item", "Qty", "Price"],
                    ["Site visit", "1", "$120"],
                    ["Labour", "2", "$180"],
                ],
            )
            # The header row is rendered bold.
            header_runs = [
                run for cell in table.rows[0].cells
                for paragraph in cell.paragraphs for run in paragraph.runs
            ]
            self.assertTrue(header_runs and all(run.bold for run in header_runs))

    def test_write_table_without_header_rows_only(self):
        """A header-less table block (rows only) still renders a valid grid."""
        from docx import Document
        with tempfile.TemporaryDirectory() as d:
            out = Path(d) / "rt_table_nohdr.docx"
            blocks = [
                {"type": "table", "rows": [["A", "1"], ["B", "2"]]},
            ]
            proc = _run_wrapper(
                "write_docx.py", ["--out", str(out), "--blocks", json.dumps(blocks)]
            )
            self.assertEqual(proc.returncode, 0, proc.stderr)
            doc = Document(str(out))
            self.assertEqual(len(doc.tables), 1)
            grid = [[c.text for c in row.cells] for row in doc.tables[0].rows]
            self.assertEqual(grid, [["A", "1"], ["B", "2"]])


@unittest.skipUnless(HAS_REPORTLAB and HAS_PDFPLUMBER,
                     "reportlab and pdfplumber both needed for the PDF table round-trip")
class TestPdfTableExtraction(unittest.TestCase):
    """Generate a small PDF with a real table grid via reportlab, then extract it
    back with pdf_tables.py and assert the rows come through."""

    def _make_table_pdf(self, path: Path, table_rows: list[list[str]]) -> None:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        from reportlab.lib import colors

        doc = SimpleDocTemplate(str(path), pagesize=A4)
        tbl = Table(table_rows)
        tbl.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ]))
        doc.build([tbl])

    def test_extract_table_from_generated_pdf(self):
        with tempfile.TemporaryDirectory() as d:
            pdf = Path(d) / "fixture.pdf"
            table_rows = [["Item", "Price"], ["Widget", "10"], ["Gadget", "25"]]
            self._make_table_pdf(pdf, table_rows)
            self.assertTrue(pdf.is_file())

            proc = _run_wrapper("pdf_tables.py", [str(pdf)])
            self.assertEqual(proc.returncode, 0, proc.stderr)
            data = json.loads(proc.stdout)
            self.assertIn("pages", data)
            # flatten every extracted table cell into one blob to assert content
            all_cells = []
            for page in data["pages"]:
                for table in page.get("tables", []):
                    for row in table:
                        all_cells.extend(c for c in row if c)
            for expected in ("Item", "Price", "Widget", "10", "Gadget", "25"):
                self.assertIn(expected, all_cells)

    def test_extract_text_mode_from_generated_pdf(self):
        with tempfile.TemporaryDirectory() as d:
            pdf = Path(d) / "text.pdf"
            # a plain paragraph PDF via the production make_pdf wrapper
            blocks = [{"type": "paragraph", "text": "Hello extraction world"}]
            mk = _run_wrapper("make_pdf.py", ["--out", str(pdf), "--blocks", json.dumps(blocks)])
            self.assertEqual(mk.returncode, 0, mk.stderr)

            proc = _run_wrapper("pdf_tables.py", [str(pdf), "--text"])
            self.assertEqual(proc.returncode, 0, proc.stderr)
            data = json.loads(proc.stdout)
            joined = " ".join(p.get("text", "") for p in data["pages"])
            self.assertIn("Hello extraction world", joined)


if __name__ == "__main__":
    unittest.main()
